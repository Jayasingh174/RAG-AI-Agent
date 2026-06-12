from docx import Document
import logging
import os

logger = logging.getLogger(__name__)

def extract_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file, ensuring that both standard paragraphs
    and embedded tables (crucial for RFQs and BOQs) are captured.
    """
    if not os.path.exists(file_path):
        logger.error(f"❌ DOCX not found at path: {file_path}")
        return ""

    try:
        doc = Document(file_path)
        content_chunks = []

        # 1️⃣ Extract standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_chunks.append(para.text.strip())

        # 2️⃣ Extract tables (CRITICAL FOR RFQs)
        if doc.tables:
            content_chunks.append("\n--- DOCUMENT TABLES ---")
            for table_idx, table in enumerate(doc.tables):
                content_chunks.append(f"\n[Table {table_idx + 1}]")
                
                for row in table.rows:
                    # Extract text from each cell, clean up internal line breaks, and separate with pipes
                    row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    
                    # Only append the row if it actually contains text
                    if any(row_data):
                        content_chunks.append(" | ".join(row_data))

        logger.info(f"✅ Successfully extracted text and {len(doc.tables)} tables from DOCX.")
        return "\n".join(content_chunks)

    except Exception as e:
        logger.error(f"❌ Critical failure extracting DOCX {file_path}. Error: {e}")
        raise ValueError(f"Failed to read DOCX file: {e}")